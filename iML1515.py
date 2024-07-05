import pandas as pd
import cobra
from cobra.io import read_sbml_model
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import warnings

# Load the E. coli iJO1366 COBRA model from the XML file
model_path = '/content/iJO1366.xml'  # Update with your path to iJO1366.xml
model = read_sbml_model(model_path)

# Define significant reactions for E. coli iJO1366 model
significant_reactions = [
    'BIOMASS_Ec_iJO1366_WT_53p95M', 'ATPM', 'PGI', 'PFK', 'PYK', 'TPI', 'GAPD', 'PGK', 'PGM', 'ENO', 'PYR', 'LDH_D',
    'G6PDH2r', 'GND', 'PGL', 'RPI', 'RPE', 'TKT1', 'TKT2', 'TALA', 'ACONTa', 'ACONTb', 'AKGDH', 'CS', 'ICDHyr', 'SUCDi', 'SUCOAS'
]

# Define key pathways and associated reactions for E. coli iJO1366 model
key_pathways = {
    'Glycolysis': ['PGI', 'PFK', 'TPI', 'GAPD', 'PGK', 'PGM', 'ENO', 'PYK'],
    'TCA Cycle': ['ACONTa', 'ACONTb', 'AKGDH', 'CS', 'ICDHyr', 'SUCDi', 'SUCOAS']
}

# Set up a range of conditions for E. coli model
conditions = {
    'Aerobic High O2': {'EX_o2_e': (-20, 1000)},
    'Aerobic Medium O2': {'EX_o2_e': (-10, 1000)},
    'Aerobic Low O2': {'EX_o2_e': (-5, 1000)},
    'Anaerobic': {'EX_o2_e': (0, 0)},
    'High Glucose': {'EX_glc__D_e': (-20, 1000)},
    'Medium Glucose': {'EX_glc__D_e': (-10, 1000)},
    'Low Glucose': {'EX_glc__D_e': (-5, 1000)},
    'No Glucose': {'EX_glc__D_e': (0, 0)},
    'High Glycerol': {'EX_glyc_e': (-20, 1000)},
    'Medium Glycerol': {'EX_glyc_e': (-10, 1000)},
    'Low Glycerol': {'EX_glyc_e': (-5, 1000)},
    'No Glycerol': {'EX_glyc_e': (0, 0)},
    'High Acetate': {'EX_ac_e': (-20, 1000)},
    'Medium Acetate': {'EX_ac_e': (-10, 1000)},
    'Low Acetate': {'EX_ac_e': (-5, 1000)},
    'No Acetate': {'EX_ac_e': (0, 0)},
    'High Ammonium': {'EX_nh4_e': (-20, 1000)},
    'Medium Ammonium': {'EX_nh4_e': (-10, 1000)},
    'Low Ammonium': {'EX_nh4_e': (-5, 1000)},
    'No Ammonium': {'EX_nh4_e': (0, 0)}
}

# Perform FBA under different conditions and store results
results = {}

for condition, bounds in conditions.items():
    with model:
        # Set the objective function to E. coli biomass reaction
        model.objective = 'BIOMASS_Ec_iJO1366_WT_53p95M'

        for reaction_id, (lower_bound, upper_bound) in bounds.items():
            model.reactions.get_by_id(reaction_id).bounds = (lower_bound, upper_bound)

        solution = model.optimize()

        # Collect results and additional debug information
        if solution.status == 'optimal':
            biomass_flux = solution.fluxes.get('BIOMASS_Ec_iJO1366_WT_53p95M', 0.0)
            growth_rate = solution.objective_value

            results[condition] = {
                'growth_rate': growth_rate,
                'biomass_flux': biomass_flux,
                'reaction_fluxes': {rxn: solution.fluxes.get(rxn, 0.0) for rxn in significant_reactions},
                'pathway_fluxes': {pathway: {rxn: solution.fluxes.get(rxn, 0.0) for rxn in rxns} for pathway, rxns in key_pathways.items()}
            }
        else:
            warnings.warn(f"Solver status is '{solution.status}'. Condition '{condition}' is infeasible.", UserWarning)

# Print and save results for debugging
output_file = "ecoli_iJO1366_fba_results.txt"
with open(output_file, 'w') as f:
    for condition, result in results.items():
        f.write(f"Condition: {condition}\n")
        f.write(f"Growth rate: {result['growth_rate']}\n")
        f.write(f"Biomass flux: {result['biomass_flux']}\n")
        for rxn, flux in result['reaction_fluxes'].items():
            f.write(f"{rxn}: {flux}\n")
        f.write("\n")
print(f"Debug results saved to '{output_file}'")

# Visualize all reaction fluxes for feasible conditions and save plots
for condition, result in results.items():
    flux_data = []
    for reaction, flux in result['reaction_fluxes'].items():
        flux_data.append({
            'Condition': condition,
            'Reaction': reaction,
            'Flux': flux
        })
    flux_df = pd.DataFrame(flux_data)

    plt.figure(figsize=(12, 6))
    sns.barplot(data=flux_df, x='Condition', y='Flux', hue='Reaction', palette='viridis', dodge=False, legend=False)
    plt.title(f'Reaction Fluxes under {condition}')
    plt.xticks(rotation=90)
    plt.tight_layout()
    plt.savefig(f'reaction_fluxes_{condition}_iJO1366.png')
    plt.close()

# Generate a detailed report
doc = Document()
doc.add_heading('Metabolic Pathway Analysis Report (E. coli - iJO1366)', level=1)

for condition, result in results.items():
    doc.add_heading(condition, level=2)
    doc.add_paragraph(f'Growth rate: {result["growth_rate"]}\n')
    doc.add_paragraph(f'Biomass flux: {result["biomass_flux"]}\n')

    doc.add_heading('Significant Reaction Fluxes', level=3)
    for rxn, flux in result['reaction_fluxes'].items():
        doc.add_paragraph(f'{rxn}: {flux}')

    doc.add_heading('Key Pathway Fluxes', level=3)
    for pathway, fluxes in result['pathway_fluxes'].items():
        doc.add_heading(pathway, level=4)
        for rxn, flux in fluxes.items():
            doc.add_paragraph(f'{rxn}: {flux}')

    # Add plots for each condition
    doc.add_heading('Visualizations', level=3)
    doc.add_picture(f'reaction_fluxes_{condition}_iJO1366.png', width=Inches(6.0))

# Save the report to a file
report_path = 'ecoli_metabolic_pathway_analysis_report_iJO1366.docx'
doc.save(report_path)
print(f"Metabolic pathway analysis report for E. coli iJO1366 saved to '{report_path}'")
